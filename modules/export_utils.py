import io
import re
from datetime import datetime

import pandas as pd
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from docx import Document


def _verdict_color(verdict: str) -> str:
    v = (verdict or "").strip().upper()
    if "ПРОДВИГАТЬ" in v:
        return "C6EFCE"
    if "ДОРАБОТАТЬ" in v:
        return "FFEB9C"
    if "ОТКЛОНИТЬ" in v:
        return "FFC7CE"
    return "FFFFFF"


def export_swot_xlsx(df: pd.DataFrame) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        total = len(df)
        verdict_col_name = None
        for c in df.columns:
            if "вердикт" in str(c).lower():
                verdict_col_name = c
                break
        if verdict_col_name:
            v_series = df[verdict_col_name].fillna("").astype(str).str.upper()
            go_count = int((v_series.str.contains("ПРОДВИГАТЬ")).sum())
            fix_count = int((v_series.str.contains("ДОРАБОТАТЬ")).sum())
            stop_count = int((v_series.str.contains("ОТКЛОНИТЬ")).sum())
            summary_data = [
                ["Показатель", "Значение"],
                ["Всего предложений", total],
                ["Продвигать", go_count],
                ["Доработать", fix_count],
                ["Отклонить", stop_count],
            ]
        else:
            summary_data = [["Показатель", "Значение"], ["Всего предложений", total]]
        pd.DataFrame(summary_data).to_excel(writer, sheet_name="Сводка", index=False, header=False)
        summary_sheet = writer.sheets["Сводка"]
        for col in range(1, summary_sheet.max_column + 1):
            summary_sheet.cell(1, col).font = Font(bold=True, color="FFFFFF")
            summary_sheet.cell(1, col).fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
        df.to_excel(writer, sheet_name="SWOT-анализ", index=False)
        sheet = writer.sheets["SWOT-анализ"]
        header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        for col in range(1, sheet.max_column + 1):
            cell = sheet.cell(1, col)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
        verdict_col = None
        for c, name in enumerate(df.columns, start=1):
            if "вердикт" in str(name).lower():
                verdict_col = c
                break
        if verdict_col:
            for row in range(2, sheet.max_row + 1):
                cell = sheet.cell(row, verdict_col)
                fill = PatternFill(
                    start_color=_verdict_color(cell.value or ""),
                    end_color=_verdict_color(cell.value or ""),
                    fill_type="solid",
                )
                cell.fill = fill
        for col in range(1, sheet.max_column + 1):
            sheet.column_dimensions[get_column_letter(col)].width = 20
        for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
    buffer.seek(0)
    return buffer.getvalue()


def export_swot_csv(df: pd.DataFrame) -> str:
    return df.to_csv(index=False, encoding="utf-8-sig")


def export_swot_md(df: pd.DataFrame) -> str:
    def _row_to_md(row):
        return "| " + " | ".join(str(x) if pd.notna(x) else "" for x in row) + " |"
    lines = [_row_to_md(df.columns), "|" + "|".join("---" for _ in df.columns) + "|"]
    for _, r in df.iterrows():
        lines.append(_row_to_md(r))
    return "\n".join(lines)


def export_swot_docx(df: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading("Отчёт по SWOT-анализу предложений", 0)
    doc.add_paragraph(f"Дата выгрузки: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")
    doc.add_heading("Сводка", level=1)
    total = len(df)
    verdict_col_name = None
    for c in df.columns:
        if "вердикт" in str(c).lower():
            verdict_col_name = c
            break
    if verdict_col_name:
        v_series = df[verdict_col_name].fillna("").astype(str).str.upper()
        go_count = (v_series.str.contains("ПРОДВИГАТЬ")).sum()
        fix_count = (v_series.str.contains("ДОРАБОТАТЬ")).sum()
        stop_count = (v_series.str.contains("ОТКЛОНИТЬ")).sum()
        doc.add_paragraph(f"Всего предложений: {total}.")
        doc.add_paragraph(f"Продвигать: {go_count}. Доработать: {fix_count}. Отклонить: {stop_count}.")
    else:
        doc.add_paragraph(f"Всего предложений: {total}.")
    doc.add_paragraph("")
    doc.add_heading("Детальный анализ", level=1)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr[i].text = str(col_name)
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
    for row_idx, (_, r) in enumerate(df.iterrows(), start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, val in enumerate(r):
            row_cells[col_idx].text = str(val) if pd.notna(val) else ""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_chat_txt(messages: list) -> str:
    lines = []
    for m in messages:
        role = "Вы" if m.get("role") == "user" else "Аналитик"
        lines.append(f"{role}:\n{m.get('content', '')}\n")
    return "\n".join(lines)


def export_chat_md(messages: list) -> str:
    lines = ["# Экспорт диалога\n", f"*{datetime.now().strftime('%d.%m.%Y %H:%M')}*\n"]
    for m in messages:
        role = "**Вы**" if m.get("role") == "user" else "**Аналитик**"
        content = (m.get("content") or "").replace("\n", "\n\n")
        lines.append(f"{role}\n\n{content}\n\n---\n")
    return "\n".join(lines)


def export_chat_docx(messages: list) -> bytes:
    doc = Document()
    doc.add_heading("Экспорт диалога с консультантом", 0)
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")
    for m in messages:
        role = "Вы" if m.get("role") == "user" else "Аналитик"
        p = doc.add_paragraph()
        p.add_run(role + ": ").bold = True
        p.add_run(m.get("content") or "")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _strip_markdown(s: str) -> str:
    if not s:
        return ""
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"__([^_]+)__", r"\1", s)
    s = re.sub(r"_([^_]+)_", r"\1", s)
    return s.strip()


def parse_swot_response(thesis: str, raw_response: str) -> dict:
    result = {
        "Тезис": thesis[:500],
        "Эффект": "",
        "Риски": "",
        "Вердикт": "",
        "Статус": "Готово",
    }
    text = (raw_response or "").strip()
    for block in text.split("---"):
        block = block.strip()
        if not block:
            continue
        for line in block.split("\n"):
            line = line.strip()
            if line.upper().startswith("ТЕЗИС:"):
                result["Тезис"] = _strip_markdown(line[6:].strip() or thesis[:500])
            elif line.upper().startswith("ЭФФЕКТ:"):
                result["Эффект"] = _strip_markdown(line[7:].strip())
            elif line.upper().startswith("РИСКИ:"):
                result["Риски"] = _strip_markdown(line[6:].strip())
            elif line.upper().startswith("ВЕРДИКТ:"):
                result["Вердикт"] = _strip_markdown(line[8:].strip())
    if not result["Вердикт"] and text:
        verdict_match = re.search(
            r"(Продвигать|Доработать|Отклонить)[\s\-—:]*([^\n]*)",
            text,
            re.IGNORECASE,
        )
        if verdict_match:
            result["Вердикт"] = _strip_markdown(
                f"{verdict_match.group(1)} - {verdict_match.group(2).strip()}"
            )
    return result
